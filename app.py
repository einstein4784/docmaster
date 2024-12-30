from flask import Flask, request, render_template, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
import os
from docx import Document
from PyPDF2 import PdfReader
import mammoth
import pdfplumber

app = Flask(__name__)
app.secret_key = 'supersecretkey'  # Required for flash messages

# Configuration
UPLOAD_FOLDER = 'uploads'
HTML_FOLDER = 'html_files'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(HTML_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Routes
@app.route('/')
def index():
    """Render the homepage with an upload form."""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file uploads and conversion."""
    if 'file' not in request.files:
        flash('No file part')
        print("No file part in the request.")  # Debugging
        return redirect(url_for('index'))

    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        print("No file selected.")  # Debugging
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    print(f"File saved at: {file_path}")  # Debugging

    # Determine the file type and convert
    html_path = os.path.join(HTML_FOLDER, f"{os.path.splitext(filename)[0]}.html")
    try:
        if filename.endswith('.docx'):
            print(f"Converting DOCX file: {file_path}")  # Debugging
            convert_docx_to_html(file_path, html_path)
        elif filename.endswith('.pdf'):
            print(f"Converting PDF file: {file_path}")  # Debugging
            convert_pdf_to_html(file_path, html_path)
        elif filename.endswith('.rtf'):
            print(f"Converting RTF file: {file_path}")  # Debugging
            convert_rtf_to_html(file_path, html_path)
        else:
            flash('Unsupported file format')
            print(f"Unsupported file format: {filename}")  # Debugging
            return redirect(url_for('index'))
    except Exception as e:
        flash(f'Error processing file: {str(e)}')
        print(f"Error during file processing: {e}")  # Debugging
        return redirect(url_for('index'))

    flash('File converted successfully!')
    print(f"HTML file generated at: {html_path}")  # Debugging
    return send_file(html_path, as_attachment=True)

# Conversion Functions
from docx import Document
import os

def convert_docx_to_html(input_path, output_path, output_dir="output_images"):
    print(f"Starting conversion for: {input_path}")  # Debugging
    os.makedirs(output_dir, exist_ok=True)
    
    try:
        doc = Document(input_path)
        print("Document loaded successfully.")  # Debugging
        
        html = """
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; padding: 10px; }
                table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
                td, th { border: 1px solid black; padding: 5px; text-align: left; vertical-align: top; }
                .bold { font-weight: bold; }
                .italic { font-style: italic; }
            </style>
        </head>
        <body>
        """
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                html += f"<p>"
                for run in paragraph.runs:
                    classes = []
                    if run.bold:
                        classes.append("bold")
                    if run.italic:
                        classes.append("italic")
                    style = " ".join(classes)
                    html += f"<span class='{style}'>{run.text}</span>"
                html += "</p>"
        print("Paragraphs processed.")  # Debugging
        
        for table in doc.tables:
            html += "<table>"
            for row in table.rows:
                html += "<tr>"
                for cell in row.cells:
                    html += f"<td>{cell.text}</td>"
                html += "</tr>"
            html += "</table>"
        print("Tables processed.")  # Debugging

        html += "</body></html>"
        with open(output_path, "w") as f:
            f.write(html)
        print(f"HTML written to: {output_path}")  # Debugging

    except Exception as e:
        print(f"Error during conversion: {e}")
        raise


def convert_pdf_to_html(input_path, output_path):
    """Convert a PDF to HTML using PDFPlumber."""
    html = """
    <html>
    <head>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; padding: 10px; }
            table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
            td, th { border: 1px solid black; padding: 5px; text-align: left; vertical-align: top; }
            .checkbox { display: inline-block; width: 16px; height: 16px; border: 1px solid black; margin-right: 5px; }
            label { display: inline-block; margin-right: 15px; font-size: 14px; vertical-align: middle; }
        </style>
    </head>
    <body>
    """

    with pdfplumber.open(input_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            html += f"<h2>Page {page_number}</h2>"

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                html += "<table>"
                for row in table:
                    html += "<tr>"
                    for cell in row:
                        cell = cell or ""
                        # Detect and replace checkbox placeholders
                        if "[ ]" in cell:
                            options = cell.split("[ ]")
                            cell = "".join(f"<label><input type='checkbox' class='checkbox'> {opt.strip()}</label>" for opt in options if opt.strip())
                        html += f"<td>{cell}</td>"
                    html += "</tr>"
                html += "</table>"

            # Extract raw text
            text = page.extract_text()
            if text:
                html += f"<pre>{text}</pre>"

    html += "</body></html>"

    with open(output_path, 'w') as f:
        f.write(html)

import pypandoc

def convert_rtf_to_html(input_path, output_path):
    print(f"Starting RTF to HTML conversion for: {input_path}")  # Debugging
    try:
        output = pypandoc.convert_file(input_path, 'html', format='rtf')
        print("RTF conversion successful.")  # Debugging
        with open(output_path, 'w') as f:
            f.write(output)
        print(f"RTF file converted to HTML: {output_path}")  # Debugging
    except Exception as e:
        print(f"Error converting RTF file: {e}")
        raise


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
